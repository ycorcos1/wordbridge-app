from __future__ import annotations

import io

import pandas as pd
from docx import Document

from app.services import text_extraction


def test_extract_text_from_txt():
    data = b"Hello world\nThis is a test."
    text = text_extraction.extract_text(data, "sample.txt")
    assert text == "Hello world This is a test."
    assert text_extraction.word_count(text) == 6


def test_extract_text_from_docx():
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")

    buffer = io.BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()

    text = text_extraction.extract_text(docx_bytes, "writing.docx")
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_text_from_csv():
    frame = pd.DataFrame({"col1": ["alpha", "beta"], "col2": ["gamma", "delta"]})

    buffer = io.BytesIO()
    frame.to_csv(buffer, index=False)
    csv_bytes = buffer.getvalue()

    text = text_extraction.extract_text(csv_bytes, "words.csv")
    assert "alpha" in text
    assert "delta" in text


def test_extract_text_rejects_unknown_extension():
    data = b"irrelevant"
    try:
        text_extraction.extract_text(data, "archive.zip")
    except text_extraction.UnsupportedFileTypeError as exc:
        assert "Unsupported" in str(exc)
    else:  # pragma: no cover - ensure exception raised
        raise AssertionError("UnsupportedFileTypeError not raised")

